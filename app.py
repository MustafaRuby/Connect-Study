from flask import Flask, render_template, request, jsonify
from flask_socketio import SocketIO, emit
import os
import uuid
import re
import traceback
from datetime import datetime
from docx import Document
import PyPDF2
from werkzeug.utils import secure_filename
import db_manager

app = Flask(__name__)
app.config['SECRET_KEY'] = 'secret!'
socketio = SocketIO(app)

# Filtri personalizzati per Jinja2
@app.template_filter('format_datetime')
def format_datetime(date_string):
    """Formatta una stringa di data in formato italiano"""
    if not date_string:
        return 'Data non disponibile'
    try:
        # Parsing della data dal formato SQLite: '2025-06-07 16:36:16'
        from datetime import datetime
        dt = datetime.strptime(date_string, '%Y-%m-%d %H:%M:%S')
        return dt.strftime('%d/%m/%Y %H:%M')
    except:
        return date_string

def is_bullet_point(paragraph):
    """Determina se un paragrafo è un bullet point basandosi su vari criteri"""
    
    # Criterio 1: Nome dello stile contiene 'Bullet'
    if 'Bullet' in paragraph.style.name:
        return True
    
    # Criterio 2: Il testo inizia con simboli comuni di bullet
    text = paragraph.text.strip()
    bullet_symbols = ['•', '●', '○', '◦', '▪', '▫', '■', '□', '‣', '⁃', '-', '*']
    
    for symbol in bullet_symbols:
        if text.startswith(symbol + ' ') or text.startswith(symbol + '\t'):
            return True
    
    # Criterio 3: Controllo delle proprietà di formattazione (numbering)
    try:
        # Accesso alle proprietà di numerazione del paragrafo
        if hasattr(paragraph, '_element'):
            # Cerca elementi di numerazione
            numbering = paragraph._element.xpath('.//w:numPr')
            if numbering:
                # Se ha proprietà di numerazione, controlla se è bullet
                for num_pr in numbering:
                    # Cerca il tipo di numero
                    num_fmt = num_pr.xpath('.//w:numFmt[@w:val]')
                    if num_fmt:
                        # Prova a ottenere il valore senza namespace specifico
                        val = None
                        for attr_name in num_fmt[0].attrib:
                            if attr_name.endswith('val'):
                                val = num_fmt[0].attrib[attr_name]
                                break
                        
                        # Formati bullet comuni in Word
                        bullet_formats = ['bullet', 'none', 'symbol']
                        if val in bullet_formats:
                            return True
                    
                    # Criterio aggiuntivo: Se non ha formato numFmt ma è 'List Paragraph', probabilmente è bullet
                    if not num_fmt and paragraph.style.name == 'List Paragraph':
                        return True
            
            # Criterio aggiuntivo: Cerca se ha proprietà di tab con allineamento bullet
            tabs = paragraph._element.xpath('.//w:tabs/w:tab[@w:val="left"]')
            if tabs and any(symbol in text for symbol in ['•', '●', '○', '◦', '▪', '▫', '■', '□', '‣', '⁃']):
                return True
                
    except (AttributeError, Exception):
        # Se non riusciamo ad accedere alle proprietà, continuiamo con altri metodi
        pass
    
    # Criterio 4: Lista con stile che contiene termini relativi ai bullet
    style_name_lower = paragraph.style.name.lower()
    bullet_keywords = ['bullet', 'point', 'item', 'unordered', 'list paragraph']
    if any(keyword in style_name_lower for keyword in bullet_keywords):
        # Per 'List Paragraph', aggiungiamo un controllo più specifico
        if 'list paragraph' in style_name_lower:
            # Se è List Paragraph e non inizia con un numero, probabilmente è bullet
            if not re.match(r'^\d+\.', text.strip()):
                return True
        else:
            return True
    
    # Criterio 5: Pattern più aggressivo - se il testo sembra una lista puntata
    if text:
        # Controlla se il paragrafo inizia come un tipico elemento di lista
        list_indicators = [
            r'^[•●○◦▪▫■□‣⁃]\s+',  # Simboli bullet
            r'^[-*+]\s+',          # Trattini e asterischi
            r'^\s*[•●○◦▪▫■□‣⁃]\s+', # Simboli bullet con spazio iniziale
            r'^\s*[-*+]\s+',       # Trattini con spazio iniziale
        ]
        
        for pattern in list_indicators:
            if re.match(pattern, text):
                return True
        
        # Criterio aggiuntivo: se il paragrafo è corto e senza punteggiatura finale,
        # potrebbe essere un elemento di lista non formattato correttamente
        if (len(text) < 200 and 
            not text.endswith(('.', '!', '?', ':')) and 
            len(text.split()) > 2 and
            paragraph.style.name in ['List Paragraph', 'ListParagraph', 'Normal']):
            
            # Se il testo contiene parole chiave tipiche di liste
            list_words = ['primo', 'secondo', 'terzo', 'inoltre', 'infine', 'anche', 'pure']
            text_lower = text.lower()
            if any(word in text_lower for word in list_words):
                return True
    
    return False

def clean_existing_bullets(text):
    """Rimuove simboli di bullet esistenti dall'inizio del testo"""
    
    # Lista di simboli bullet comuni
    bullet_symbols = ['•', '●', '○', '◦', '▪', '▫', '■', '□', '‣', '⁃', '-', '*']
    
    text = text.strip()
    
    # Rimuovi bullet symbols all'inizio seguiti da spazio o tab
    for symbol in bullet_symbols:
        # Pattern che cerca il simbolo all'inizio seguito da spazio/tab
        pattern = r'^' + re.escape(symbol) + r'[\s\t]+'
        text = re.sub(pattern, '', text)
    
    # Rimuovi anche bullet symbols all'inizio senza spazio
    for symbol in bullet_symbols:
        if text.startswith(symbol):
            text = text[1:].lstrip()
            break
    
    return text

def convert_docx_to_markdown(file_path):
    """Converte un file .docx in markdown con formattazione avanzata"""
    try:
        doc = Document(file_path)
        
        # Prima passata: analizza tutte le dimensioni dei font per creare una mappa dei livelli
        font_sizes = analyze_font_sizes(doc)
        
        markdown_content = []
        
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                markdown_content.append('')
                continue
                
            # Gestione dei titoli basata sullo stile
            if paragraph.style.name.startswith('Heading'):
                level = int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1
                markdown_content.append('#' * level + ' ' + paragraph.text.strip())
                markdown_content.append('')
                continue
            # Gestione delle liste - Prima controlla se è un bullet point
            # Questo controllo ora viene prima per tutti i paragrafi che potrebbero essere liste
            is_bullet = is_bullet_point(paragraph)
            
            if paragraph.style.name.startswith('List') or is_bullet:
                if is_bullet:
                    # È un bullet point
                    processed_text = process_paragraph_formatting(paragraph, font_sizes)
                    # Rimuovi bullet symbol esistenti se presenti
                    cleaned_text = clean_existing_bullets(processed_text)
                    markdown_content.append('- ' + cleaned_text)
                else:
                    # È una lista ma non un bullet point, quindi lista numerata
                    formatted_text = process_paragraph_formatting(paragraph, font_sizes)
                    if formatted_text.strip().startswith('#'):
                        # È un titolo, non una lista
                        markdown_content.append(formatted_text)
                    else:
                        # Controlla se il testo ha già numerazione
                        if re.match(r'^\d+\.', paragraph.text.strip()):
                            # Ha già numerazione, mantienila
                            markdown_content.append(formatted_text)
                        else:
                            # È una lista numerata senza numerazione esplicita
                            markdown_content.append('1. ' + formatted_text)
                continue
            
            # Gestione dei paragrafi normali con formattazione
            formatted_text = process_paragraph_formatting(paragraph, font_sizes)
            if formatted_text.strip():
                markdown_content.append(formatted_text)
                markdown_content.append('')
          # Gestione delle tabelle
        for table in doc.tables:
            markdown_content.append(convert_table_to_markdown(table))
            markdown_content.append('')
        
        return '\n'.join(markdown_content).strip()
    except Exception as e:
        # Logging dell'errore completo per debugging
        print(f"Errore dettagliato nella conversione Word: {type(e).__name__}: {str(e)}")
        traceback.print_exc()
        
        # Prova a fare una conversione di base se quella avanzata fallisce
        try:
            doc = Document(file_path)
            basic_content = []
            basic_content.append("# Documento convertito (modalità semplificata)")
            basic_content.append("")
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Solo testo semplice, senza formattazione
                    basic_content.append(text)
                    basic_content.append("")
            
            return '\n'.join(basic_content).strip()
        except Exception as fallback_error:
            return f"Errore nella conversione del file Word: {str(e)}\n\nErrore anche nella conversione semplificata: {str(fallback_error)}"

def analyze_font_sizes(doc):
    """Analizza tutte le dimensioni dei font nel documento per creare una mappa dei livelli"""
    font_sizes = set()
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            try:
                if run.font.size and run.font.size.pt:
                    font_sizes.add(run.font.size.pt)
            except (AttributeError, TypeError):
                # Ignora run senza informazioni di font valide
                continue
    
    # Ordina le dimensioni in ordine decrescente
    sorted_sizes = sorted(font_sizes, reverse=True)
    
    # Crea una mappa: dimensione -> livello di titolo
    size_to_level = {}
    
    if len(sorted_sizes) <= 1:
        # Se c'è solo una dimensione o nessuna, non creare titoli automatici
        return size_to_level
    
    # Assegna livelli solo se ci sono significative differenze di dimensione
    for i, size in enumerate(sorted_sizes):
        if i == 0:
            # La dimensione più grande diventa h1 solo se significativamente più grande
            if len(sorted_sizes) > 1 and size - sorted_sizes[1] >= 4:
                size_to_level[size] = 1
        elif i == 1:
            # La seconda più grande diventa h2 solo se significativamente più grande del testo normale
            if len(sorted_sizes) > 2 and size - sorted_sizes[-1] >= 2:
                size_to_level[size] = 2
        elif i == 2:
            # La terza più grande diventa h3 solo se significativamente più grande del testo normale
            if size - sorted_sizes[-1] >= 1:
                size_to_level[size] = 3
    
    return size_to_level

def process_paragraph_formatting(paragraph, font_sizes_map):
    """Processa la formattazione di un paragrafo"""
    result = []
    paragraph_has_heading = False
    full_text = paragraph.text.strip()
    
    # Controlla se questo paragrafo è un titolo numerato (es: "1. Titolo", "2. Titolo")
    is_numbered_heading = is_numbered_title(full_text, font_sizes_map, paragraph)
    
    # Controlla se questo paragrafo dovrebbe essere un titolo basato sulla dimensione del font
    for run in paragraph.runs:
        if run.font.size and run.font.size.pt in font_sizes_map:
            level = font_sizes_map[run.font.size.pt]
            # Se questo run ha una dimensione che corrisponde a un titolo
            if not paragraph_has_heading:
                if is_numbered_heading:
                    # Per titoli numerati, usa il formato: ## 1. Titolo
                    result.append('#' * level + ' ')
                else:
                    result.append('#' * level + ' ')
                paragraph_has_heading = True
            break
    
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
              # Applica formattazione basata sulle proprietà del run con gestione errori
        try:
            formatted_text = apply_run_formatting(run, text, font_sizes_map, paragraph_has_heading)
            result.append(formatted_text)
        except Exception as format_error:
            # Se la formattazione fallisce, usa il testo semplice
            print(f"Errore nella formattazione del run: {format_error}")
            result.append(text)
    
    return ''.join(result)

def is_numbered_title(text, font_sizes_map, paragraph):
    """Determina se un paragrafo è un titolo numerato"""
    # Controlla se inizia con un numero seguito da punto e spazio
      # Pattern per titoli numerati: "1. ", "2. ", "3.1 ", etc.
    numbered_pattern = r'^\d+(\.\d+)*\.\s+'
    
    if re.match(numbered_pattern, text):
        # Controlla se ha una dimensione di font che indica un titolo
        for run in paragraph.runs:
            try:
                if run.font.size and run.font.size.pt in font_sizes_map:
                    return True
            except (AttributeError, TypeError):
                continue
        
        # Anche se non ha font size specifico, se il testo è corto e sembra un titolo
        remaining_text = re.sub(numbered_pattern, '', text).strip()
        if len(remaining_text) < 100 and len(remaining_text.split()) <= 15:
            return True
    
    return False

def apply_run_formatting(run, text, font_sizes_map, is_heading_paragraph):
    """Applica la formattazione markdown basata sulle proprietà del run"""
    if not text.strip():
        return text
    
    formatted = text
    
    # Non applicare formattazione di titolo se siamo già in un paragrafo di titolo
    # o se la dimensione del font non è significativamente diversa
    apply_font_size_formatting = not is_heading_paragraph
      # Grassetto - gestione sicura delle proprietà
    try:
        if getattr(run, 'bold', False):
            formatted = f"**{formatted}**"
    except (AttributeError, TypeError):
        pass
    
    # Corsivo
    try:
        if getattr(run, 'italic', False):
            formatted = f"*{formatted}*"
    except (AttributeError, TypeError):
        pass
    
    # Sottolineato (usando HTML per compatibilità)
    try:
        if getattr(run, 'underline', False):
            formatted = f"<u>{formatted}</u>"
    except (AttributeError, TypeError):
        pass
    
    # Barrato
    try:
        if getattr(run.font, 'strike', False):
            formatted = f"~~{formatted}~~"
    except (AttributeError, TypeError):
        pass
      # Apice (superscript)
    try:
        if getattr(run.font, 'superscript', False):
            formatted = f"<sup>{formatted}</sup>"
    except (AttributeError, TypeError):
        pass
    
    # Pedice (subscript)
    try:
        if getattr(run.font, 'subscript', False):
            formatted = f"<sub>{formatted}</sub>"
    except (AttributeError, TypeError):
        pass
      # Colore del testo (se diverso dal nero)
    try:
        if run.font.color.rgb:
            rgb = run.font.color.rgb
            # python-docx RGBColor ha attributi r, g, b (non red, green, blue)
            if hasattr(rgb, 'r') and hasattr(rgb, 'g') and hasattr(rgb, 'b'):
                # Controlla se il colore è diverso dal nero
                if (rgb.r, rgb.g, rgb.b) != (0, 0, 0):
                    color_hex = f"#{rgb.r:02x}{rgb.g:02x}{rgb.b:02x}"
                    formatted = f'<span style="color: {color_hex};">{formatted}</span>'
            elif hasattr(rgb, '__iter__') and len(rgb) >= 3:
                # Fallback per tuple/list RGB
                r, g, b = rgb[0], rgb[1], rgb[2]
                if (r, g, b) != (0, 0, 0):
                    color_hex = f"#{r:02x}{g:02x}{b:02x}"
                    formatted = f'<span style="color: {color_hex};">{formatted}</span>'
    except (AttributeError, TypeError) as e:
        # Ignora errori di colore per non bloccare la conversione
        pass
      # Dimensione del font - ora gestita in modo intelligente
    try:
        if apply_font_size_formatting and run.font.size and run.font.size.pt:
            size = run.font.size.pt
            # Solo per testo molto piccolo, usa <small>
            if size < 9:
                formatted = f"<small>{formatted}</small>"
    except (AttributeError, TypeError):
        pass
        # Non applicare più automaticamente ### per dimensioni medie
      # Evidenziazione (highlight)
    try:
        if run.font.highlight_color and run.font.highlight_color != 0:
            formatted = f"=={formatted}=="
    except (AttributeError, TypeError):
        # Ignora errori di highlight per non bloccare la conversione
        pass
      # Font famiglia (se diversa dalla standard)
    try:
        if run.font.name and run.font.name.lower() in ['courier new', 'consolas', 'monaco', 'monospace']:
            formatted = f"`{formatted}`"
    except (AttributeError, TypeError):
        pass
    
    return formatted

def convert_table_to_markdown(table):
    """Converte una tabella Word in formato Markdown"""
    markdown_rows = []
    
    for i, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            cell_text = cell.text.strip().replace('\n', ' ')
            cells.append(cell_text)
        
        # Riga della tabella
        markdown_rows.append('| ' + ' | '.join(cells) + ' |')
        
        # Separatore dopo la prima riga (header)
        if i == 0:
            separator = '| ' + ' | '.join(['---'] * len(cells)) + ' |'
            markdown_rows.append(separator)
    
    return '\n'.join(markdown_rows)

def convert_pdf_to_markdown(file_path):
    """Converte un file .pdf in markdown con miglior gestione del testo"""
    try:
        markdown_content = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():
                    markdown_content.append(f"## Pagina {page_num + 1}")
                    markdown_content.append('')
                    
                    # Processamento del testo per migliorare la formattazione
                    processed_text = process_pdf_text(text)
                    markdown_content.append(processed_text)
                    markdown_content.append('')
        
        return '\n'.join(markdown_content).strip()
    except Exception as e:
        return f"Errore nella conversione del file PDF: {str(e)}"

def process_pdf_text(text):
    """Processa il testo estratto dal PDF per migliorare la formattazione"""
    lines = text.split('\n')
    processed_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Riconosce titoli numerati (es: "1. Introduzione", "2.1 Metodi")
        numbered_title_pattern = r'^\d+(\.\d+)*\.\s+\w+'
        if re.match(numbered_title_pattern, line) and len(line) < 100:
            # Determina il livello basandosi sulla numerazione
            number_part = re.match(r'^(\d+(\.\d+)*)', line).group(1)
            dots_count = number_part.count('.')
            level = min(dots_count + 1, 4)  # Massimo h4
            processed_lines.append('#' * level + ' ' + line)
            continue
            
        # Riconosce possibili titoli (linee più corte in maiuscolo)
        if len(line) < 60 and line.isupper() and len(line.split()) <= 6:
            processed_lines.append(f"### {line.title()}")
        # Riconosce possibili sottotitoli (iniziano con maiuscola e sono più corti)
        elif len(line) < 80 and line[0].isupper() and line.count('.') == 0 and len(line.split()) <= 10:
            # Controlla se potrebbe essere un titolo
            words = line.split()
            if all(word[0].isupper() for word in words if len(word) > 2):
                processed_lines.append(f"#### {line}")
            else:
                processed_lines.append(line)
        # Riconosce liste (iniziano con numero o simbolo) ma non titoli numerati
        elif line.startswith(('•', '-', '*', '◦')) or (line[0].isdigit() and line[1:3] in ['. ', ') '] and not re.match(numbered_title_pattern, line)):
            if not line.startswith(('- ', '* ')):
                if line.startswith(('•', '◦')):
                    line = '- ' + line[1:].strip()
                elif line[0].isdigit() and not re.match(numbered_title_pattern, line):
                    line = '1. ' + line[2:].strip()
            processed_lines.append(line)
        else:
            processed_lines.append(line)
    
    return '\n\n'.join(processed_lines)

def convert_file_to_markdown(file_path, filename):
    """Converte un file in markdown basandosi sulla sua estensione"""
    ext = filename.lower().split('.')[-1]
    
    if ext == 'docx':
        return convert_docx_to_markdown(file_path)
    elif ext == 'doc':
        # Per i file .doc, prova prima a convertirli (richiede python-docx2txt)
        try:
            import docx2txt
            text = docx2txt.process(file_path)
            return f"# Documento convertito da {filename}\n\n{text}"
        except ImportError:
            return "Formato .doc non supportato. Installa docx2txt o converti in .docx"
    elif ext == 'pdf':
        return convert_pdf_to_markdown(file_path)
    elif ext in ['md', 'markdown']:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception as e:
                return f"Errore nella lettura del file Markdown: {str(e)}"
    elif ext == 'txt':
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                # Converte il testo semplice in markdown basico
                return convert_txt_to_markdown(content, filename)
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
                    return convert_txt_to_markdown(content, filename)
            except Exception as e:
                return f"Errore nella lettura del file di testo: {str(e)}"
    elif ext in ['rtf']:
        return "Formato RTF non ancora supportato. Converti in .docx o .txt"
    else:
        return f"Formato file .{ext} non supportato. Usa .docx, .pdf, .md, .txt"

def convert_txt_to_markdown(content, filename):
    """Converte un file di testo semplice in markdown con formattazione basica"""
    lines = content.split('\n')
    markdown_lines = [f"# {filename.replace('.txt', '')}", ""]
    
    current_paragraph = []
    
    for line in lines:
        line = line.strip()
        
        # Riga vuota - termina il paragrafo corrente
        if not line:
            if current_paragraph:
                markdown_lines.append(' '.join(current_paragraph))
                markdown_lines.append('')
                current_paragraph = []
            continue
        
        # Possibile titolo (riga corta, maiuscola, poche parole)
        if (len(line) < 50 and 
            line.isupper() and 
            len(line.split()) <= 6 and 
            not any(char in line for char in '.,;:!?')):
            if current_paragraph:
                markdown_lines.append(' '.join(current_paragraph))
                markdown_lines.append('')
                current_paragraph = []
            markdown_lines.append(f"## {line.title()}")
            markdown_lines.append('')
          # Lista numerata normale (non titolo)
        elif line.startswith(tuple(f"{i}." for i in range(1, 100))) and len(line.split()) > 10:
            if current_paragraph:
                markdown_lines.append(' '.join(current_paragraph))
                markdown_lines.append('')
                current_paragraph = []
            markdown_lines.append(line)
        
        # Titolo numerato (es: "1. Introduzione", "2. Metodologia")
        elif line.startswith(tuple(f"{i}." for i in range(1, 100))) and len(line.split()) <= 10:
            if current_paragraph:
                markdown_lines.append(' '.join(current_paragraph))
                markdown_lines.append('')
                current_paragraph = []
            # Determina il livello dal numero di punti nella numerazione
            number_match = re.match(r'^(\d+(\.\d+)*)', line)
            if number_match:
                number_part = number_match.group(1)
                level = min(number_part.count('.') + 2, 4)  # h2 minimo, h4 massimo
                markdown_lines.append('#' * level + ' ' + line)
            else:
                markdown_lines.append(f"## {line}")
            markdown_lines.append('')
        
        # Lista puntata
        elif line.startswith(('-', '*', '•')):
            if current_paragraph:
                markdown_lines.append(' '.join(current_paragraph))
                markdown_lines.append('')
                current_paragraph = []
            if not line.startswith('- '):
                line = '- ' + line[1:].strip()
            markdown_lines.append(line)
        
        # Testo normale
        else:
            current_paragraph.append(line)
    
    # Aggiungi l'ultimo paragrafo se presente
    if current_paragraph:
        markdown_lines.append(' '.join(current_paragraph))
    
    return '\n'.join(markdown_lines).strip()

# Inizializzazione DB se non esiste
db_manager.init_database()

@app.route('/')
def index():
    materie = db_manager.get_all_materie()
    return render_template('index.html', materie=materie)

@app.route('/add_materia', methods=['POST'])
def add_materia():
    nome = request.form['nome']
    colore = request.form.get('colore', '#cccccc')
    db_manager.add_materia(nome, colore)
    socketio.emit('update_materie')
    return ('', 204)

@app.route('/edit_materia/<int:id>', methods=['POST'])
def edit_materia(id):
    nome = request.form['nome']
    colore = request.form.get('colore', '#cccccc')
    db_manager.update_materia(id, nome, colore)
    socketio.emit('update_materie')
    return ('', 204)

@app.route('/delete_materia/<int:id>', methods=['DELETE'])
def delete_materia(id):
    db_manager.delete_materia(id)
    socketio.emit('update_materie')
    return ('', 204)

@socketio.on('reorder_materie')
def reorder_materie(data):
    # data: lista di id materie nell'ordine nuovo
    db_manager.reorder_materie(data['order'])
    emit('update_materie', broadcast=True)

@app.route('/api/materie')
def api_materie():
    materie = db_manager.get_all_materie()
    return jsonify([dict(m) for m in materie])

@app.route('/api/argomenti/materia/<int:id_materia>')
def api_argomenti_by_materia(id_materia):
    """API endpoint per ottenere argomenti di una materia specifica"""
    argomenti = db_manager.get_argomenti_by_materia(id_materia)
    return jsonify([dict(a) for a in argomenti])

@app.route('/materia/<int:id>')
def materia_argomenti(id):
    materia = db_manager.get_materia_by_id(id)
    if not materia:
        return "Materia non trovata", 404
    argomenti = db_manager.get_argomenti_by_materia(id)
    return render_template('argomenti.html', materia=materia, argomenti=argomenti)

@app.route('/add_argomento', methods=['POST'])
def add_argomento():
    id_materia = request.form['id_materia']
    titolo = request.form['titolo']
    contenuto_md = request.form.get('contenuto_md', '')
    colore = request.form.get('colore', '#cccccc')
    etichetta_preparazione = request.form.get('etichetta_preparazione', 'scarsa preparazione')
    
    # Gestione file caricato
    if 'file_content' in request.files:
        file = request.files['file_content']
        if file.filename != '' and file.filename is not None:
            # Salva temporaneamente il file
            filename = secure_filename(file.filename)
            temp_path = os.path.join(os.path.dirname(__file__), 'temp_' + filename)
            file.save(temp_path)
            
            try:
                # Converte il file in markdown
                file_markdown = convert_file_to_markdown(temp_path, filename)
                
                # Gestisce la modalità di inserimento
                file_mode = request.form.get('file_mode', 'replace')
                if file_mode == 'replace':
                    contenuto_md = file_markdown
                elif file_mode == 'prepend':
                    contenuto_md = file_markdown + '\n\n' + contenuto_md if contenuto_md else file_markdown
                elif file_mode == 'append':
                    contenuto_md = contenuto_md + '\n\n' + file_markdown if contenuto_md else file_markdown
                
            finally:
                # Rimuove il file temporaneo
                if os.path.exists(temp_path):
                    os.remove(temp_path)
    
    db_manager.add_argomento(id_materia, titolo, contenuto_md, colore, etichetta_preparazione)
    socketio.emit('update_argomenti', {'id_materia': id_materia})
    return ('', 204)

@app.route('/api/argomenti/<int:id_materia>')
def api_argomenti(id_materia):
    argomenti = db_manager.get_argomenti_by_materia(id_materia)
    return jsonify([dict(a) for a in argomenti])

# Route per il singolo argomento
@app.route('/argomento/<int:id>')
def argomento_dettaglio(id):
    argomento = db_manager.get_argomento_by_id(id)
    if not argomento:
        return "Argomento non trovato", 404
    
    materia = db_manager.get_materia_by_id(argomento['id_materia'])
    allegati = db_manager.get_allegati_by_argomento(id)
    
    return render_template('argomento.html', argomento=argomento, materia=materia, allegati=allegati)

# Route per aggiornare il contenuto dell'argomento
@app.route('/update_argomento/<int:id>', methods=['POST'])
def update_argomento(id):
    contenuto_md = request.form.get('contenuto_md', '')
    titolo = request.form.get('titolo')
    colore = request.form.get('colore')
    etichetta_preparazione = request.form.get('etichetta_preparazione')
    
    # Gestione file caricato per aggiornamento
    if 'file_content' in request.files:
        file = request.files['file_content']
        if file.filename != '' and file.filename is not None:
            filename = secure_filename(file.filename)
            temp_path = os.path.join(os.path.dirname(__file__), 'temp_' + filename)
            file.save(temp_path)
            
            try:
                file_markdown = convert_file_to_markdown(temp_path, filename)
                
                file_mode = request.form.get('file_mode', 'replace')
                if file_mode == 'replace':
                    contenuto_md = file_markdown
                elif file_mode == 'prepend':
                    contenuto_md = file_markdown + '\n\n' + contenuto_md if contenuto_md else file_markdown
                elif file_mode == 'append':
                    contenuto_md = contenuto_md + '\n\n' + file_markdown if contenuto_md else file_markdown
            finally:
                if os.path.exists(temp_path):
                    os.remove(temp_path)
    
    # Determina quale funzione usare in base ai campi forniti
    if titolo and colore:
        # Aggiornamento completo (titolo, colore, contenuto, etichetta)
        db_manager.update_argomento_full(id, titolo, colore, contenuto_md, etichetta_preparazione)
    elif etichetta_preparazione:
        # Aggiorna contenuto e etichetta
        db_manager.update_argomento_content_and_label(id, contenuto_md, etichetta_preparazione)
    else:
        # Aggiorna solo contenuto
        db_manager.update_argomento_content(id, contenuto_md)
    
    socketio.emit('update_argomento', {'id': id})
    return ('', 204)

# Route per caricare allegati
@app.route('/add_allegato/<int:id_argomento>', methods=['POST'])
def add_allegato(id_argomento):
    if 'file' not in request.files:
        return 'Nessun file caricato', 400
    
    file = request.files['file']
    if file.filename == '':
        return 'Nessun file selezionato', 400
    
    if file:
        # Crea la directory se non esiste
        upload_dir = os.path.join(app.static_folder, 'argomenti')
        os.makedirs(upload_dir, exist_ok=True)
          # Genera un nome file sicuro e univoco
        original_filename = secure_filename(file.filename)
        base_name, ext = os.path.splitext(original_filename)
        
        # Usa timestamp e UUID per garantire unicità
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        unique_id = str(uuid.uuid4())[:8]
        filename = f"{base_name}_{timestamp}_{unique_id}{ext}"
        
        file_path = os.path.join(upload_dir, filename)
        file.save(file_path)
          # Salva nel database
        db_manager.add_allegato(id_argomento, filename, f'static/argomenti/{filename}')
        
        socketio.emit('update_allegati', {'id_argomento': id_argomento})
        return ('', 204)

# Route per eliminare allegati
@app.route('/delete_allegato/<int:id>', methods=['DELETE'])
def delete_allegato(id):
    allegato = db_manager.get_allegato_by_id(id)
    if allegato:
        # Rimuovi il file fisico
        file_path = os.path.join(os.path.dirname(__file__), allegato['percorso'])
        if os.path.exists(file_path):
            os.remove(file_path)
        
        # Rimuovi dal database
        db_manager.delete_allegato(id)
        id_argomento = allegato['id_argomento']
        
        socketio.emit('update_allegati', {'id_argomento': id_argomento})
        return ('', 204)
    
    return ('File non trovato', 404)

# Route per eliminare un argomento
@app.route('/delete_argomento/<int:id>', methods=['DELETE'])
def delete_argomento(id):
    argomento = db_manager.get_argomento_by_id(id)
    if not argomento:
        return 'Argomento non trovato', 404
    
    # Elimina tutti gli allegati associati
    allegati = db_manager.get_allegati_by_argomento(id)
    for allegato in allegati:
        file_path = os.path.join(os.path.dirname(__file__), allegato['percorso'])
        if os.path.exists(file_path):
            os.remove(file_path)
        db_manager.delete_allegato(allegato['id'])
    
    # Elimina l'argomento
    db_manager.delete_argomento(id)
    socketio.emit('update_argomenti', {'id_materia': argomento['id_materia']})
    return ('', 204)

# API per ottenere allegati di un argomento
@app.route('/api/allegati/<int:id_argomento>')
def api_allegati(id_argomento):
    allegati = db_manager.get_allegati_by_argomento(id_argomento)
    return jsonify([dict(a) for a in allegati])

# === COLLEGAMENTI ROUTES ===

@app.route('/collegamenti')
def collegamenti_list():
    """Pagina con tutti i collegamenti"""
    collegamenti = db_manager.get_all_collegamenti()
    materie = db_manager.get_all_materie()
    return render_template('collegamenti.html', collegamenti=collegamenti, materie=materie)

@app.route('/collegamenti/materia/<int:id_materia>')
def collegamenti_materia(id_materia):
    """Pagina con i collegamenti di una materia specifica"""
    materia = db_manager.get_materia_by_id(id_materia)
    if not materia:
        return "Materia non trovata", 404
    collegamenti = db_manager.get_collegamenti_by_materia(id_materia)
    materie = db_manager.get_all_materie()
    return render_template('collegamenti.html', collegamenti=collegamenti, materie=materie, 
                         materia_selezionata=materia)

@app.route('/add_collegamento', methods=['POST'])
def add_collegamento():
    """Aggiunge un nuovo collegamento"""
    try:
        titolo = request.form['titolo']
        id_argomento1 = request.form['id_argomento1']
        id_argomento2 = request.form['id_argomento2']
        dettagli = request.form.get('dettagli', '')
        etichetta_qualita = request.form.get('etichetta_qualita', 'collegamento media qualità')
        
        # Verifica che gli argomenti siano diversi
        if id_argomento1 == id_argomento2:
            return jsonify({'success': False, 'error': 'Non puoi collegare un argomento a se stesso'})
        
        db_manager.add_collegamento(titolo, id_argomento1, id_argomento2, dettagli, etichetta_qualita)
        socketio.emit('update_collegamenti')
        return jsonify({'success': True, 'message': 'Collegamento creato con successo'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/update_collegamento/<int:id>', methods=['POST'])
def update_collegamento(id):
    """Aggiorna un collegamento esistente"""
    try:
        titolo = request.form['titolo']
        dettagli = request.form.get('dettagli', '')
        etichetta_qualita = request.form.get('etichetta_qualita', 'collegamento media qualità')
        
        # Raccoglie anche gli argomenti se forniti
        id_argomento1 = request.form.get('id_argomento1')
        id_argomento2 = request.form.get('id_argomento2')
        
        # Verifica che gli argomenti siano diversi se entrambi sono forniti
        if id_argomento1 and id_argomento2 and id_argomento1 == id_argomento2:
            return jsonify({'success': False, 'error': 'Non puoi collegare un argomento a se stesso'})
        
        # Se entrambi gli argomenti sono forniti, aggiorna tutto
        if id_argomento1 and id_argomento2:
            db_manager.update_collegamento(id, titolo, dettagli, etichetta_qualita, 
                                         int(id_argomento1), int(id_argomento2))
        else:
            # Altrimenti aggiorna solo titolo, dettagli ed etichetta
            db_manager.update_collegamento(id, titolo, dettagli, etichetta_qualita)
        
        socketio.emit('update_collegamenti')
        return jsonify({'success': True, 'message': 'Collegamento aggiornato con successo'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/delete_collegamento/<int:id>', methods=['DELETE'])
def delete_collegamento(id):
    """Elimina un collegamento"""
    db_manager.delete_collegamento(id)
    socketio.emit('update_collegamenti')
    return ('', 204)

# API per ottenere collegamenti
@app.route('/api/collegamenti')
def api_collegamenti():
    """API per ottenere tutti i collegamenti"""
    collegamenti = db_manager.get_all_collegamenti()
    return jsonify([dict(c) for c in collegamenti])

@app.route('/api/collegamenti/materia/<int:id_materia>')
def api_collegamenti_materia(id_materia):
    """API per ottenere i collegamenti di una materia"""
    collegamenti = db_manager.get_collegamenti_by_materia(id_materia)
    return jsonify([dict(c) for c in collegamenti])

@app.route('/api/collegamenti/argomento/<int:id_argomento>')
def api_collegamenti_argomento(id_argomento):
    """API per ottenere i collegamenti di un argomento"""
    collegamenti = db_manager.get_collegamenti_by_argomento(id_argomento)
    return jsonify([dict(c) for c in collegamenti])

@app.route('/api/search_collegamenti')
def api_search_collegamenti():
    """API per cercare collegamenti"""
    query_titolo = request.args.get('titolo', '')
    query_dettagli = request.args.get('dettagli', '')
    etichetta_qualita = request.args.get('etichetta_qualita', '')
    query_argomenti = request.args.get('argomenti', '')
    query_materia = request.args.get('materia', '')
    
    collegamenti = db_manager.search_collegamenti(query_titolo, query_dettagli, etichetta_qualita, query_argomenti, query_materia)
    return jsonify([dict(c) for c in collegamenti])

# === SIMULAZIONI ROUTES ===

@app.route('/simulazioni')
def simulazioni():
    """Pagina delle simulazioni"""
    simulazioni = db_manager.get_all_simulazioni()
    edit_id = request.args.get('edit')  # Get edit parameter from URL
    return render_template('simulazioni.html', simulazioni=simulazioni, edit_id=edit_id)

@app.route('/add_simulazione', methods=['POST'])
def add_simulazione():
    """Aggiunge una nuova simulazione"""
    try:
        spunto_testo = request.form.get('spunto_testo', '')
        spunto_immagine = None
        
        # Gestione upload immagine
        if 'spunto_immagine' in request.files:
            file = request.files['spunto_immagine']
            if file and file.filename:
                # Crea la cartella se non esiste
                upload_folder = os.path.join('static', 'simulazioni')
                os.makedirs(upload_folder, exist_ok=True)
                  # Nome univoco per il file
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                unique_id = str(uuid.uuid4().hex[:8])
                filename = f"spunto_{timestamp}_{unique_id}.{file.filename.rsplit('.', 1)[1].lower()}"
                # Usa '/' per URL web invece di os.path.join che usa '\' su Windows
                spunto_immagine = f"simulazioni/{filename}"
                
                file.save(os.path.join('static', spunto_immagine))
        
        # Verifica che almeno uno dei due sia presente
        if not spunto_testo and not spunto_immagine:
            return jsonify({'success': False, 'error': 'Inserisci almeno un testo o un\'immagine come spunto'})
        
        db_manager.add_simulazione(spunto_testo, spunto_immagine)
        socketio.emit('update_simulazioni')
        return jsonify({'success': True, 'message': 'Simulazione creata con successo'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/simulazione/<int:id>')
def simulazione_detail(id):
    """Pagina dettaglio di una simulazione"""
    simulazione = db_manager.get_simulazione_by_id(id)
    if not simulazione:
        return "Simulazione non trovata", 404
    
    fili = db_manager.get_fili_by_simulazione(id)
    materie = db_manager.get_all_materie()
    return render_template('simulazione_detail.html', simulazione=simulazione, fili=fili, materie=materie)

@app.route('/delete_simulazione/<int:id>', methods=['DELETE'])
def delete_simulazione(id):
    """Elimina una simulazione esistente"""
    try:
        # Get current simulation to remove image file
        current_sim = db_manager.get_simulazione_by_id(id)
        if current_sim and current_sim['spunto_immagine']:
            image_path = os.path.join('static', current_sim['spunto_immagine'])
            if os.path.exists(image_path):
                os.remove(image_path)
        
        # Delete the simulation
        db_manager.delete_simulazione(id)
        socketio.emit('update_simulazioni')
        return jsonify({'success': True, 'message': 'Simulazione eliminata con successo'})
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/simulazione/<int:id>')
def api_simulazione_by_id(id):
    """API endpoint per ottenere una simulazione specifica"""
    try:
        simulazione = db_manager.get_simulazione_by_id(id)
        if simulazione:
            return jsonify(dict(simulazione))
        else:
            return jsonify({'error': 'Simulazione non trovata'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/update_simulazione/<int:id>', methods=['POST'])
def update_simulazione(id):
    """Aggiorna una simulazione esistente"""
    try:
        spunto_testo = request.form.get('spunto_testo', '')
        current_sim = db_manager.get_simulazione_by_id(id)
        
        if not current_sim:
            return jsonify({'success': False, 'error': 'Simulazione non trovata'})
        
        spunto_immagine = current_sim['spunto_immagine']  # Keep current image by default
        
        # Handle image upload
        if 'spunto_immagine' in request.files:
            file = request.files['spunto_immagine']
            if file and file.filename:
                # Remove old image if exists
                if spunto_immagine:
                    old_image_path = os.path.join('static', spunto_immagine)
                    if os.path.exists(old_image_path):
                        os.remove(old_image_path)
                
                # Save new image
                upload_folder = os.path.join('static', 'simulazioni')
                os.makedirs(upload_folder, exist_ok=True)
                
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                unique_id = str(uuid.uuid4().hex[:8])
                filename = f"spunto_{timestamp}_{unique_id}.{file.filename.rsplit('.', 1)[1].lower()}"
                spunto_immagine = f"simulazioni/{filename}"
                
                file.save(os.path.join('static', spunto_immagine))
        
        # Verify at least one spunto is present
        if not spunto_testo and not spunto_immagine:
            return jsonify({'success': False, 'error': 'Inserisci almeno un testo o un\'immagine come spunto'})
        
        # Update simulation
        db_manager.update_simulazione(id, spunto_testo, spunto_immagine)
        socketio.emit('update_simulazioni')
        return jsonify({'success': True, 'message': 'Simulazione aggiornata con successo'})
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

# === FILI COLLEGAMENTO ROUTES ===

@app.route('/add_filo_simulazione', methods=['POST'])
def add_filo_simulazione():
    """Aggiunge un nuovo filo di collegamento a una simulazione"""
    try:
        id_simulazione = request.form['id_simulazione']
        db_manager.add_filo_collegamento(id_simulazione)
        socketio.emit('update_fili', {'id_simulazione': id_simulazione})
        return jsonify({'success': True, 'message': 'Filo creato con successo'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/delete_filo/<int:id>', methods=['DELETE'])
def delete_filo(id):
    """Elimina un filo di collegamento e tutti i suoi collegamenti"""
    try:
        filo = db_manager.get_filo_by_id(id)
        
        if filo:
            # Ottieni l'ID della simulazione dal filo - converti in dict per accesso sicuro
            filo_dict = dict(filo)
            id_simulazione = filo_dict.get('id_simulazione')
            
            db_manager.delete_filo_collegamento(id)
            
            if id_simulazione:
                socketio.emit('update_fili', {'id_simulazione': id_simulazione})
        else:
            return jsonify({'success': False, 'error': 'Filo non trovato'}), 404
        
        return ('', 204)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

# === COLLEGAMENTI SIMULAZIONE ROUTES ===

@app.route('/add_collegamento_simulazione', methods=['POST'])
def add_collegamento_simulazione():
    """Aggiunge un nuovo collegamento di simulazione"""
    try:
        id_filo = request.form['id_filo']
        titolo = request.form['titolo']
        id_argomento2 = request.form['id_argomento2']
        id_argomento1 = request.form.get('id_argomento1') or None
        dettagli = request.form.get('dettagli', '')
        etichetta_qualita = request.form.get('etichetta_qualita', 'collegamento media qualità')
        
        # Verifica che gli argomenti siano diversi se entrambi sono forniti
        if id_argomento1 and id_argomento1 == id_argomento2:
            return jsonify({'success': False, 'error': 'Non puoi collegare un argomento a se stesso'})
        
        db_manager.add_collegamento_simulazione(id_filo, titolo, id_argomento2, id_argomento1, dettagli, etichetta_qualita)
        socketio.emit('update_collegamenti_simulazione', {'id_filo': id_filo})
        return jsonify({'success': True, 'message': 'Collegamento creato con successo'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/import_collegamento_to_simulazione', methods=['POST'])
def import_collegamento_to_simulazione():
    """Importa un collegamento esistente in un filo di simulazione"""
    try:
        collegamento_id = request.form['collegamento_id']
        id_filo = request.form['id_filo']
        
        db_manager.copy_collegamento_to_simulazione(collegamento_id, id_filo)
        socketio.emit('update_collegamenti_simulazione', {'id_filo': id_filo})
        return jsonify({'success': True, 'message': 'Collegamento importato con successo'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

# === API ROUTES PER SIMULAZIONI ===

@app.route('/api/collegamenti_simulazione/filo/<int:id_filo>')
def api_collegamenti_simulazione_filo(id_filo):
    """API per ottenere i collegamenti di simulazione di un filo"""
    collegamenti = db_manager.get_collegamenti_simulazione_by_filo(id_filo)
    return jsonify([dict(c) for c in collegamenti])

@app.route('/api/collegamenti_simulazione/<int:collegamento_id>')
def api_collegamento_simulazione_detail(collegamento_id):
    """API per ottenere i dettagli di un singolo collegamento di simulazione"""
    collegamento = db_manager.get_collegamento_simulazione_by_id(collegamento_id)
    if collegamento:
        return jsonify(dict(collegamento))
    else:
        return jsonify({'error': 'Collegamento non trovato'}), 404

@app.route('/api/collegamento_simulazione/<int:id>', methods=['GET'])
def get_collegamento_simulazione(id):
    """Ottiene i dati di un collegamento di simulazione"""
    try:
        collegamento = db_manager.get_collegamento_simulazione_by_id(id)
        if collegamento:
            return jsonify(collegamento)
        else:
            return jsonify({'error': 'Collegamento non trovato'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/update_collegamento_simulazione/<int:id>', methods=['POST'])
def update_collegamento_simulazione(id):
    """Aggiorna un collegamento di simulazione esistente"""
    try:
        titolo = request.form['titolo']
        dettagli = request.form.get('dettagli', '')
        etichetta_qualita = request.form.get('etichetta_qualita', 'collegamento media qualità')
          # Raccoglie anche gli argomenti se forniti
        id_argomento1 = request.form.get('id_argomento1')
        id_argomento2 = request.form.get('id_argomento2')
        
        # Converte valori vuoti in None
        id_argomento1 = int(id_argomento1) if id_argomento1 and id_argomento1.strip() else None
        id_argomento2 = int(id_argomento2) if id_argomento2 and id_argomento2.strip() else None
        
        # Verifica che almeno un argomento sia fornito
        if not id_argomento2:
            return jsonify({'success': False, 'error': 'Il secondo argomento è obbligatorio'})
        
        # Verifica che gli argomenti siano diversi se entrambi sono forniti
        if id_argomento1 and id_argomento2 and id_argomento1 == id_argomento2:
            return jsonify({'success': False, 'error': 'Non puoi collegare un argomento a se stesso'})
          # Aggiorna sempre tutti i campi, inclusi gli argomenti
        db_manager.update_collegamento_simulazione(id, titolo, dettagli, etichetta_qualita, 
                                     id_argomento1, id_argomento2)
        
        # Get the filo_id to emit the correct event
        collegamento = db_manager.get_collegamento_simulazione_by_id(id)
        if collegamento:
            socketio.emit('update_collegamenti_simulazione', {'id_filo': collegamento['id_filo']})
        
        return jsonify({'success': True, 'message': 'Collegamento aggiornato con successo'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/delete_collegamento_simulazione/<int:id>', methods=['DELETE'])
def delete_collegamento_simulazione(id):
    """Elimina un collegamento di simulazione"""
    try:
        # Get the filo_id before deleting to emit the correct event
        collegamento = db_manager.get_collegamento_simulazione_by_id(id)
        if collegamento:
            id_filo = collegamento['id_filo']
            db_manager.delete_collegamento_simulazione(id)
            socketio.emit('update_collegamenti_simulazione', {'id_filo': id_filo})
            return ('', 204)
        else:
            return jsonify({'error': 'Collegamento non trovato'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/reorder_collegamento_simulazione/<int:id>', methods=['POST'])
def reorder_collegamento_simulazione(id):
    """Cambia l'ordine di un collegamento di simulazione"""
    try:
        new_order = int(request.form.get('new_order', 0))
        if new_order <= 0:
            return jsonify({'success': False, 'error': 'Numero d\'ordine non valido'})
        
        # Get the collegamento before reordering to get filo info
        collegamento = db_manager.get_collegamento_simulazione_by_id(id)
        if not collegamento:
            return jsonify({'success': False, 'error': 'Collegamento non trovato'})
        
        id_filo = collegamento['id_filo']
        max_order = db_manager.get_max_order_in_filo(id_filo)
        
        # Verifica che il nuovo ordine sia valido
        if new_order > max_order:
            return jsonify({'success': False, 'error': f'Ordine massimo consentito: {max_order}'})
        
        success = db_manager.reorder_collegamento_simulazione(id, new_order)
        
        if success:
            socketio.emit('update_collegamenti_simulazione', {'id_filo': id_filo})
            return jsonify({'success': True, 'message': 'Ordine aggiornato con successo'})
        else:
            return jsonify({'success': False, 'error': 'Errore durante il riordino'})
            
    except ValueError:
        return jsonify({'success': False, 'error': 'Numero d\'ordine non valido'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/search_argomenti')
def api_search_argomenti():
    """API per cercare argomenti"""
    query = request.args.get('q', '')
    if len(query) < 2:
        return jsonify([])
    
    argomenti = db_manager.search_argomenti(query)
    return jsonify([dict(a) for a in argomenti])

if __name__ == '__main__':
    # Run on all available network interfaces (LAN access)
    socketio.run(app, host='0.0.0.0', debug=True)
